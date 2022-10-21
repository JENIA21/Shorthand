class ReadText:

    @staticmethod
    def receiving_a_secret_message():
        file = open("copybook.txt", "r", encoding="utf-8")
        return file.read().lower()


class HidingData:
    import docx
    import sys

    index = -1
    marker = 0
    letter_in_encrypted_text = ""
    array = []

    document = docx.Document('test.docx')
    ciphertext = document.paragraphs

    for letter_in_text in ciphertext:
        letter_in_encrypted_text += letter_in_text.text + "\n"
        letter_in_encrypted_text = letter_in_encrypted_text[:-1]

    for letter in ReadText.receiving_a_secret_message():
        index = letter_in_encrypted_text.find(letter, index + 1)

        if index == -1:
            sys.exit('Not enough letters to encrypt!')
        else:
            array.append(index)

    document_wt = docx.Document()
    indent = array[0]
    text = document_wt.add_paragraph(letter_in_encrypted_text[marker:indent])

    for j in range(len(array)):
        data_for_arr = text.add_run(letter_in_encrypted_text[array[j]])
        data_for_arr.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
        if j == (len(array) - 1):
            data_for_arr = text.add_run(letter_in_encrypted_text[array[j] + 1:])
        else:
            data_for_arr = text.add_run(letter_in_encrypted_text[array[j] + 1:array[j + 1]])
    document_wt.save("test.docx")
    print("Успешно!")
